#!/usr/bin/env python3
"""
Создание технической документации для Training Tracker Power Apps
Генерирует User Guide, Deployment Guide и HLD в формате DOCX
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_user_guide():
    """Создание руководства пользователя"""

    doc = Document()
    doc.add_heading('Руководство пользователя', 0)
    doc.add_heading('Training Tracker Power Apps', 1)

    # Введение
    doc.add_heading('1. Введение', level=2)
    intro = doc.add_paragraph()
    intro.add_run('Это руководство поможет пользователям эффективно работать с приложением Training Tracker, '
                  'предназначенным для управления заявками на тренинги и курсы.').bold = True

    # Главная страница
    doc.add_heading('2. Главная страница - Статистика', level=2)
    doc.add_paragraph('Главная страница отображает общую аналитику по заявкам:')

    stats = doc.add_paragraph()
    stats.add_run('📊 Общая статистика:\n').bold = True
    stats.add_run('• Всего заявок: 24\n')
    stats.add_run('• Ожидают утверждения: 12\n')
    stats.add_run('• Утверждено: 11\n')
    stats.add_run('• Отклонено: 1\n\n')

    stats.add_run('🎯 Популярные курсы:\n').bold = True
    stats.add_run('• Python Test\n')
    stats.add_run('• Agile\n')
    stats.add_run('• SaaS\n')
    stats.add_run('(Показаны топ-3 из 15+ доступных курсов)')

    # Создание заявки
    doc.add_heading('3. Создание новой заявки', level=2)
    doc.add_paragraph('Для создания новой заявки на тренинг:')

    steps = doc.add_paragraph()
    steps.add_run('1. Нажмите кнопку "New application"\n').bold = True
    steps.add_run('2. Заполните обязательные поля:\n')
    steps.add_run('   • Title - название курса\n')
    steps.add_run('   • Location - выберите из списка (Minsk, Warsaw, Berlin, Batumi)\n')
    steps.add_run('   • Department - отдел (Power Platform, Big Data, Finance, HR)\n')
    steps.add_run('   • Type - тип (Course, Training, Exam, Certification)\n')
    steps.add_run('   • Line Manager - руководитель\n')
    steps.add_run('   • Status - статус заявки\n')
    steps.add_run('   • Start Date - дата начала (выберите в календаре)\n')
    steps.add_run('   • End Date - дата окончания (выберите в календаре)\n\n')

    steps.add_run('3. Нажмите "Submit" для отправки или "Cancel" для отмены\n').bold = True

    # Просмотр заявок
    doc.add_heading('4. Просмотр всех заявок', level=2)
    doc.add_paragraph('Таблица "All applications" показывает все заявки с возможностью фильтрации:')

    table_desc = doc.add_paragraph()
    table_desc.add_run('📋 Доступные колонки:\n').bold = True
    table_desc.add_run('• Title - название курса\n')
    table_desc.add_run('• Start Date - дата начала\n')
    table_desc.add_run('• End Date - дата окончания\n')
    table_desc.add_run('• Location - локация\n')
    table_desc.add_run('• Department - отдел\n')
    table_desc.add_run('• Type - тип тренинга\n')
    table_desc.add_run('• Line Manager - руководитель\n')
    table_desc.add_run('• Status - статус (Not Started, Approved, In Progress)\n\n')

    table_desc.add_run('🔍 Статусы заявок:\n').bold = True
    table_desc.add_run('• Not Started - не начато\n')
    table_desc.add_run('• Approved - утверждено\n')
    table_desc.add_run('• In Progress - в процессе\n')
    table_desc.add_run('• Rejected - отклонено\n')

    # Навигация
    doc.add_heading('5. Навигация и горячие клавиши', level=2)
    nav = doc.add_paragraph()
    nav.add_run('🖱️ Основные элементы интерфейса:\n').bold = True
    nav.add_run('• Кнопка "New application" - создание новой заявки\n')
    nav.add_run('• Кнопка "Cancel" - отмена действия\n')
    nav.add_run('• Кнопка "Submit" - подтверждение действия\n')
    nav.add_run('• Календарь - выбор дат\n')
    nav.add_run('• Выпадающие списки - выбор значений\n\n')

    nav.add_run('📱 Адаптивный дизайн:\n').bold = True
    nav.add_run('Приложение адаптируется под размер экрана устройства\n')
    nav.add_run('Все элементы остаются доступными на мобильных устройствах\n')

    # Устранение неполадок
    doc.add_heading('6. Устранение неполадок', level=2)
    trouble = doc.add_paragraph()
    trouble.add_run('🔧 Возможные проблемы:\n').bold = True
    trouble.add_run('• Данные не обновляются - обновите страницу (F5)\n')
    trouble.add_run('• Форма не отправляется - проверьте заполнение всех обязательных полей\n')
    trouble.add_run('• Кнопки не работают - проверьте подключение к интернету\n')
    trouble.add_run('• Даты не выбираются - используйте календарь или введите вручную\n\n')

    trouble.add_run('📞 Техническая поддержка:\n').bold = True
    trouble.add_run('При возникновении проблем обратитесь к администратору системы\n')
    trouble.add_run('или в службу технической поддержки\n')

    doc.save('User_Guide_Training_Tracker.docx')
    print('✅ Создан User_Guide_Training_Tracker.docx')

def create_deployment_guide():
    """Создание руководства по развертыванию"""

    doc = Document()
    doc.add_heading('Руководство по развертыванию', 0)
    doc.add_heading('Training Tracker Power Apps', 1)

    # Предварительные требования
    doc.add_heading('1. Предварительные требования', level=2)

    req = doc.add_paragraph()
    req.add_run('🔧 Системные требования:\n').bold = True
    req.add_run('• Microsoft 365 лицензия с Power Apps\n')
    req.add_run('• SharePoint Online доступ\n')
    req.add_run('• Современный веб-браузер (Chrome, Edge, Safari)\n')
    req.add_run('• Стабильное интернет-соединение\n\n')

    req.add_run('👤 Права доступа:\n').bold = True
    req.add_run('• Power Apps Creator лицензия\n')
    req.add_run('• Доступ к SharePoint сайту\n')
    req.add_run('• Права на создание и редактирование списков\n')
    req.add_run('• Права на публикацию приложений\n')

    # Подготовка SharePoint
    doc.add_heading('2. Подготовка SharePoint списка', level=2)

    sp_setup = doc.add_paragraph()
    sp_setup.add_run('2.1 Создание списка "Training Tracker"\n').bold = True
    sp_setup.add_run('1. Перейдите на SharePoint сайт\n')
    sp_setup.add_run('2. Нажмите "+ New" → "List"\n')
    sp_setup.add_run('3. Выберите "Blank list"\n')
    sp_setup.add_run('4. Название: "Training Tracker"\n')
    sp_setup.add_run('5. Описание: "Заявки на тренинги и курсы"\n\n')

    sp_setup.add_run('2.2 Создание колонок\n').bold = True
    sp_setup.add_run('Добавьте следующие колонки (Column settings):\n\n')

    columns_table = doc.add_table(rows=9, cols=3)
    columns_table.style = 'Table Grid'

    # Заголовки
    hdr_cells = columns_table.rows[0].cells
    hdr_cells[0].text = 'Имя колонки'
    hdr_cells[1].text = 'Тип'
    hdr_cells[2].text = 'Настройки'

    # Данные
    data = [
        ('Title', 'Single line of text', 'Обязательное поле'),
        ('Location', 'Choice', 'Minsk, Warsaw, Berlin, Batumi'),
        ('Department', 'Choice', 'Power Platform, Big Data, Finance, HR'),
        ('Type_Item', 'Choice', 'Course, Training, Exam, Certification'),
        ('Line Manager', 'Person', 'Выбор из пользователей'),
        ('Status', 'Choice', 'Not Started, Approved, In Progress, Rejected'),
        ('Start Date', 'Date and Time', 'Обязательное поле'),
        ('End Date', 'Date and Time', 'Обязательное поле'),
    ]

    for i, (name, col_type, settings) in enumerate(data, 1):
        row_cells = columns_table.rows[i].cells
        row_cells[0].text = name
        row_cells[1].text = col_type
        row_cells[2].text = settings

    # Создание Power Apps
    doc.add_heading('3. Создание Power Apps приложения', level=2)

    pa_setup = doc.add_paragraph()
    pa_setup.add_run('3.1 Создание приложения\n').bold = True
    pa_setup.add_run('1. Перейдите в Power Apps (make.powerapps.com)\n')
    pa_setup.add_run('2. Нажмите "Create" → "Blank app"\n')
    pa_setup.add_run('3. Выберите "Canvas" app\n')
    pa_setup.add_run('4. Название: "Training Tracker"\n')
    pa_setup.add_run('5. Формат: "Tablet" или "Phone"\n\n')

    pa_setup.add_run('3.2 Подключение данных\n').bold = True
    pa_setup.add_run('1. В панели "Data" нажмите "Add data"\n')
    pa_setup.add_run('2. Выберите "SharePoint"\n')
    pa_setup.add_run('3. Подключитесь к сайту с списком "Training Tracker"\n')
    pa_setup.add_run('4. Выберите список "Training Tracker"\n\n')

    pa_setup.add_run('3.3 Создание экранов\n').bold = True
    pa_setup.add_run('Создайте 3 экрана:\n')
    pa_setup.add_run('• Screen1 - Главная страница с аналитикой\n')
    pa_setup.add_run('• Screen2 - Форма создания заявки\n')
    pa_setup.add_run('• Screen3 - Галерея всех заявок\n\n')

    # Настройка формул
    doc.add_heading('4. Настройка формул и логики', level=2)

    formulas = doc.add_paragraph()
    formulas.add_run('4.1 Главная страница\n').bold = True
    formulas.add_run('Общая статистика:\n')
    formulas.add_run('CountRows(\'Training Tracker\')\n\n')

    formulas.add_run('Ожидают утверждения:\n')
    formulas.add_run('CountRows(Filter(\'Training Tracker\', Status = "Not Started"))\n\n')

    formulas.add_run('Популярные курсы:\n')
    formulas.add_run('FirstN(Distinct(\'Training Tracker\', Title), 3)\n\n')

    formulas.add_run('4.2 Форма создания\n').bold = True
    formulas.add_run('Submit button OnSelect:\n')
    formulas.add_run('SubmitForm(Form1); Navigate(Screen1)\n\n')

    formulas.add_run('4.3 Галерея заявок\n').bold = True
    formulas.add_run('Items: \'Training Tracker\'\n')
    formulas.add_run('SortByColumns: Status, Title\n\n')

    # Тестирование
    doc.add_heading('5. Тестирование и публикация', level=2)

    test = doc.add_paragraph()
    test.add_run('5.1 Тестирование функций\n').bold = True
    test.add_run('• Создайте тестовую заявку\n')
    test.add_run('• Проверьте все статусы\n')
    test.add_run('• Протестируйте на разных устройствах\n')
    test.add_run('• Проверьте производительность\n\n')

    test.add_run('5.2 Публикация\n').bold = True
    test.add_run('1. Нажмите "File" → "Save"\n')
    test.add_run('2. Нажмите "Publish"\n')
    test.add_run('3. Выберите аудиторию\n')
    test.add_run('4. Подтвердите публикацию\n\n')

    test.add_run('5.3 Мониторинг\n').bold = True
    test.add_run('• Отслеживайте использование\n')
    test.add_run('• Собирайте обратную связь\n')
    test.add_run('• Мониторьте производительность\n')

    # Безопасность
    doc.add_heading('6. Безопасность и доступ', level=2)

    security = doc.add_paragraph()
    security.add_run('6.1 Уровни доступа\n').bold = True
    security.add_run('• Администраторы: полный доступ к настройкам\n')
    security.add_run('• Пользователи: создание и просмотр заявок\n')
    security.add_run('• Руководители: утверждение заявок\n\n')

    security.add_run('6.2 SharePoint права\n').bold = True
    security.add_run('• Contribute - для создания заявок\n')
    security.add_run('• Read - для просмотра\n')
    security.add_run('• Edit - для редактирования\n\n')

    security.add_run('6.3 Аудит\n').bold = True
    security.add_run('• Включите аудит SharePoint списка\n')
    security.add_run('• Отслеживайте изменения заявок\n')
    security.add_run('• Регулярно проверяйте логи доступа\n')

    doc.save('Deployment_Guide_Training_Tracker.docx')
    print('✅ Создан Deployment_Guide_Training_Tracker.docx')

def create_hld():
    """Создание архитектуры высокого уровня"""

    doc = Document()
    doc.add_heading('Архитектура высокого уровня (HLD)', 0)
    doc.add_heading('Training Tracker Power Apps', 1)

    # Обзор архитектуры
    doc.add_heading('1. Обзор архитектуры', level=2)

    overview = doc.add_paragraph()
    overview.add_run('Training Tracker - это решение для управления заявками на корпоративные тренинги, '
                    'построенное на стеке Microsoft Power Platform.').bold = True

    overview.add_run('\n\nАрхитектура включает:\n')
    overview.add_run('• Power Apps Canvas приложение\n')
    overview.add_run('• SharePoint Online как источник данных\n')
    overview.add_run('• Microsoft 365 экосистема\n')
    overview.add_run('• Адаптивный веб-интерфейс\n')

    # Компоненты системы
    doc.add_heading('2. Компоненты системы', level=2)

    components = doc.add_paragraph()
    components.add_run('2.1 Пользовательский интерфейс\n').bold = True
    components.add_run('• Canvas Power Apps приложение\n')
    components.add_run('• Адаптивный дизайн (Tablet/Phone)\n')
    components.add_run('• Современный Material Design интерфейс\n')
    components.add_run('• Поддержка PWA (Progressive Web App)\n\n')

    components.add_run('2.2 Источник данных\n').bold = True
    components.add_run('• SharePoint Online список "Training Tracker"\n')
    components.add_run('• 8 колонок с различными типами данных\n')
    components.add_run('• Choice поля для стандартизации ввода\n')
    components.add_run('• Person поля для выбора пользователей\n')
    components.add_run('• DateTime поля для дат тренингов\n\n')

    components.add_run('2.3 Бизнес-логика\n').bold = True
    components.add_run('• Валидация форм Power Apps\n')
    components.add_run('• Расчеты статистики\n')
    components.add_run('• Фильтрация и сортировка данных\n')
    components.add_run('• Управление статусами заявок\n\n')

    components.add_run('2.4 Интеграции\n').bold = True
    components.add_run('• Microsoft 365 ecosystem\n')
    components.add_run('• SharePoint REST API\n')
    components.add_run('• Power Apps connectors\n')
    components.add_run('• Microsoft Graph API (потенциально)\n')

    # Архитектурная диаграмма
    doc.add_heading('3. Архитектурная диаграмма', level=2)

    diagram = doc.add_paragraph()
    diagram.add_run('Пользователь\n').bold = True
    diagram.add_run('    ↓\n')
    diagram.add_run('Power Apps Canvas App\n').bold = True
    diagram.add_run('    ↓ (Power Apps Connectors)\n')
    diagram.add_run('SharePoint Online\n').bold = True
    diagram.add_run('    ↓\n')
    diagram.add_run('Microsoft 365 Tenant\n').bold = True
    diagram.add_run('    ↓\n')
    diagram.add_run('Azure AD Authentication\n').bold = True

    # Поток данных
    doc.add_heading('4. Поток данных', level=2)

    data_flow = doc.add_paragraph()
    data_flow.add_run('4.1 Создание заявки:\n').bold = True
    data_flow.add_run('Пользователь → Форма Power Apps → Валидация → SharePoint List → Подтверждение\n\n')

    data_flow.add_run('4.2 Просмотр статистики:\n').bold = True
    data_flow.add_run('Пользователь → Power Apps → Запрос к SharePoint → Агрегация данных → Отображение\n\n')

    data_flow.add_run('4.3 Управление заявками:\n').bold = True
    data_flow.add_run('Пользователь → Галерея Power Apps → CRUD операции → SharePoint List → Синхронизация\n')

    # Модель данных
    doc.add_heading('5. Модель данных', level=2)

    data_model = doc.add_paragraph()
    data_model.add_run('5.1 Структура списка Training Tracker:\n').bold = True

    # Создаем таблицу с моделью данных
    data_table = doc.add_table(rows=9, cols=4)
    data_table.style = 'Table Grid'

    # Заголовки
    hdr_cells = data_table.rows[0].cells
    hdr_cells[0].text = 'Поле'
    hdr_cells[1].text = 'Тип данных'
    hdr_cells[2].text = 'Обязательное'
    hdr_cells[3].text = 'Описание'

    # Данные
    fields_data = [
        ('ID', 'Auto Number', 'Да', 'Уникальный идентификатор'),
        ('Title', 'Text', 'Да', 'Название курса/тренинга'),
        ('Location', 'Choice', 'Да', 'Место проведения'),
        ('Department', 'Choice', 'Да', 'Отдел сотрудника'),
        ('Type_Item', 'Choice', 'Да', 'Тип тренинга'),
        ('Line Manager', 'Person', 'Да', 'Руководитель'),
        ('Status', 'Choice', 'Да', 'Статус заявки'),
        ('Start Date', 'DateTime', 'Да', 'Дата начала'),
    ]

    for i, (field, data_type, required, desc) in enumerate(fields_data, 1):
        if i < len(data_table.rows):
            row_cells = data_table.rows[i].cells
            row_cells[0].text = field
            row_cells[1].text = data_type
            row_cells[2].text = required
            row_cells[3].text = desc

    # Добавим последнюю строку отдельно
    if len(data_table.rows) > 8:
        last_row = data_table.rows[8].cells
        last_row[0].text = 'End Date'
        last_row[1].text = 'DateTime'
        last_row[2].text = 'Да'
        last_row[3].text = 'Дата окончания'

    # Производительность
    doc.add_heading('6. Производительность и масштабируемость', level=2)

    perf = doc.add_paragraph()
    perf.add_run('6.1 Ограничения Power Apps:\n').bold = True
    perf.add_run('• 2000 записей для не делегированных запросов\n')
    perf.add_run('• 30 секунд таймаут для запросов\n')
    perf.add_run('• Ограничения SharePoint API\n\n')

    perf.add_run('6.2 Оптимизации производительности:\n').bold = True
    perf.add_run('• Использование делегированных запросов\n')
    perf.add_run('• Кэширование данных в коллекциях\n')
    perf.add_run('• Ленивая загрузка больших списков\n')
    perf.add_run('• Оптимизация формул Power Fx\n\n')

    perf.add_run('6.3 Масштабируемость:\n').bold = True
    perf.add_run('• Поддержка до 50,000 элементов в списке\n')
    perf.add_run('• Горизонтальное масштабирование через Power Apps\n')
    perf.add_run('• Интеграция с Power BI для аналитики\n')

    # Безопасность
    doc.add_heading('7. Безопасность', level=2)

    security = doc.add_paragraph()
    security.add_run('7.1 Аутентификация:\n').bold = True
    security.add_run('• Azure AD интеграция\n')
    security.add_run('• Single Sign-On (SSO)\n')
    security.add_run('• Multi-Factor Authentication (MFA)\n\n')

    security.add_run('7.2 Авторизация:\n').bold = True
    security.add_run('• SharePoint права доступа\n')
    security.add_run('• Row Level Security (RLS)\n')
    security.add_run('• Power Apps security roles\n\n')

    security.add_run('7.3 Шифрование:\n').bold = True
    security.add_run('• HTTPS для всех соединений\n')
    security.add_run('• Шифрование данных в SharePoint\n')
    security.add_run('• Безопасное хранение учетных данных\n')

    # Мониторинг и поддержка
    doc.add_heading('8. Мониторинг и поддержка', level=2)

    monitoring = doc.add_paragraph()
    monitoring.add_run('8.1 Логирование:\n').bold = True
    monitoring.add_run('• Power Apps Monitor\n')
    monitoring.add_run('• SharePoint audit logs\n')
    monitoring.add_run('• Azure Application Insights\n\n')

    monitoring.add_run('8.2 Метрики производительности:\n').bold = True
    monitoring.add_run('• Время загрузки экранов\n')
    monitoring.add_run('• Успешность API запросов\n')
    monitoring.add_run('• Использование памяти\n\n')

    monitoring.add_run('8.3 Резервное копирование:\n').bold = True
    monitoring.add_run('• SharePoint site backups\n')
    monitoring.add_run('• Power Apps version control\n')
    monitoring.add_run('• Disaster recovery plan\n')

    doc.save('HLD_Training_Tracker.docx')
    print('✅ Создан HLD_Training_Tracker.docx')

def main():
    """Главная функция генерации документации"""
    print('🚀 Генерация технической документации для Training Tracker...')
    print('=' * 70)

    try:
        create_user_guide()
        create_deployment_guide()
        create_hld()

        print('=' * 70)
        print('✅ Документация успешно создана!')
        print()
        print('📁 Созданные файлы:')
        print('   • User_Guide_Training_Tracker.docx')
        print('   • Deployment_Guide_Training_Tracker.docx')
        print('   • HLD_Training_Tracker.docx')
        print()
        print('📍 Расположение: /Users/user/Desktop/medical-ai-service/')

    except Exception as e:
        print(f'❌ Ошибка при создании документации: {e}')
        return 1

    return 0

if __name__ == "__main__":
    exit(main())
